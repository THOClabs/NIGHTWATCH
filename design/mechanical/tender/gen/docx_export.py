"""
Vendor DOCX exporter — assembles the authored Markdown tender documents (the
BP-00 master dossier + each bid package's SOW) into ONE Word RFQ a buyer can send.

Requires python-docx (``pip install python-docx``) — NOT a test dependency; the
Markdown under packages/ is the source of truth and the committed .docx is a
convenience deliverable. Run:  python3 -m design.mechanical.tender.gen.docx_export
"""

from __future__ import annotations

import pathlib
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

TENDER = pathlib.Path(__file__).resolve().parents[1]
PACK = TENDER / "packages"
OUT = TENDER / "docs" / "NIGHTWATCH_tender_RFQ.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
RED = RGBColor(0xA0, 0x00, 0x00)
STATUS = "ISSUED FOR BID / FOR PE REVIEW — NOT FOR CONSTRUCTION"

# Ordered assembly: the master dossier first, then each package.
DOSSIER_ORDER = [
    "README.md", "00_instructions_to_bidders.md", "01_statement_of_work_general.md",
    "02_standards_register.md", "03_drawing_register.md", "04_inspection_test_plan.md",
    "05_bid_form.md", "06_terms_and_conditions.md",
]
PACKAGE_ORDER = ["BP-01_mount_machining", "BP-02_turned_parts", "BP-03_pier_foundation",
                 "BP-04_rolloff_roof", "BP-05_cots_schedule"]
PACKAGE_FILES = ["README.md", "SOW.md", "acceptance.md", "weld_map.md", "SCHEDULE.md",
                 "fastener_schedule.md"]

_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, honouring **bold** and stripping `code` ticks."""
    text = text.replace("`", "")
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _flush_table(doc, rows: list[list[str]]) -> None:
    rows = [r for r in rows if not all(set(c.strip()) <= {"-", ":", " "} for c in r)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Light Grid Accent 1"
    for i, r in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncols):
            val = r[j] if j < len(r) else ""
            cells[j].text = ""
            _add_runs(cells[j].paragraphs[0], val)
            if i == 0:
                for run in cells[j].paragraphs[0].runs:
                    run.bold = True


def _render_markdown(doc, md: str) -> None:
    table_buf: list[list[str]] = []
    for raw in md.splitlines():
        line = raw.rstrip()
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_buf.append(cells)
            continue
        if table_buf:
            _flush_table(doc, table_buf)
            table_buf = []
        if not line.strip():
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            doc.add_heading(line.lstrip("#").strip(), level=min(level, 4))
        elif line.lstrip().startswith((">",)):
            p = doc.add_paragraph()
            r = p.add_run(line.lstrip("> ").strip())
            r.italic = True
        elif re.match(r"^\s*[-*•]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^\s*[-*•]\s+", "", line))
        elif re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\s*\d+\.\s+", "", line))
        else:
            _add_runs(doc.add_paragraph(), line)
    if table_buf:
        _flush_table(doc, table_buf)


def _title_page(doc) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("NIGHTWATCH OBSERVATORY")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Telescope Mount & Roll-off 'Turret'\nFabrication Tender / Request for Quotation")
    r2.font.size = Pt(15)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = st.add_run(STATUS)
    r3.bold = True
    r3.font.color.rgb = RED
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("Rev A · dimensions dual-unit (mm [inch]) · generated from the proven design "
                "(design/mechanical/calc/params.py)").italic = True
    doc.add_page_break()


def build() -> pathlib.Path:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _title_page(doc)

    doc.add_heading("Part I — Master Tender Dossier (BP-00)", level=1)
    dossier = PACK / "BP-00_master_dossier"
    for fname in DOSSIER_ORDER:
        f = dossier / fname
        if f.exists():
            _render_markdown(doc, f.read_text())
            doc.add_page_break()

    for pkg in PACKAGE_ORDER:
        d = PACK / pkg
        if not d.exists():
            continue
        doc.add_heading(f"Part II — {pkg.replace('_', ' ')}", level=1)
        for fname in PACKAGE_FILES:
            f = d / fname
            if f.exists():
                _render_markdown(doc, f.read_text())
        doc.add_page_break()

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print("wrote", build())
